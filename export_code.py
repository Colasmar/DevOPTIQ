import os
from docx import Document
import pydotplus as pydot

from Code.extensions import db
import Code.models.models  # pour enregistrer les modèles SQLAlchemy

# Paramètres
root_dir = os.getcwd()
excluded_dirs = {"Archives", "Backup", "venv", ".venv", ".git", "__pycache__", "Lib", "site-packages", "Scripts"}
included_extensions = (".py", ".html", ".js", ".bat", ".txt")
excluded_file_names = {"jquery-3.6.0.min.js"}

# Crée le document Word
doc = Document()
doc.add_heading("Export de l'Architecture, du Code et de la Structure de la Base de Données", level=1)

# ------------------------------------------------------------------------------
# 1. Architecture du Projet
# ------------------------------------------------------------------------------
doc.add_heading("1. Architecture du Projet", level=2)

def add_architecture(path, indent=""):
    try:
        items = sorted(os.listdir(path))
    except Exception as e:
        doc.add_paragraph(f"{indent}Erreur en listant {path}: {e}")
        return
    for item in items:
        if item.startswith("."):
            continue
        full_path = os.path.join(path, item)
        if os.path.isdir(full_path):
            if item in excluded_dirs:
                continue
            doc.add_paragraph(f"{indent}📁 {item}")
            add_architecture(full_path, indent + "    ")
        elif os.path.isfile(full_path) and item.endswith(included_extensions):
            doc.add_paragraph(f"{indent}📄 {item}")

add_architecture(root_dir)

# ------------------------------------------------------------------------------
# 2. Code Source
# ------------------------------------------------------------------------------
doc.add_heading("2. Code Source", level=2)

for folder, subdirs, files in os.walk(root_dir):
    subdirs[:] = [d for d in subdirs if d not in excluded_dirs and not d.startswith(".")]
    if any(excl in folder for excl in excluded_dirs):
        continue
    for file in files:
        if file.endswith(included_extensions):
            file_path = os.path.join(folder, file)
            doc.add_heading(f"Début du fichier: {file}", level=3)
            if file in excluded_file_names:
                doc.add_paragraph("⚠️ Fichier volumineux ou minifié ignoré pour lisibilité.")
                doc.add_heading(f"Fin du fichier: {file}", level=3)
                print(f"⏭️ Contenu ignoré : {file_path}")
                continue
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    code = f.read()
                    doc.add_paragraph(code)
                    doc.add_heading(f"Fin du fichier: {file}", level=3)
                print(f"✅ Fichier exporté : {file_path}")
            except Exception as e:
                doc.add_paragraph(f"❌ Erreur de lecture : {e}")
                print(f"❌ Erreur pour : {file_path}")

# ------------------------------------------------------------------------------
# 3. Structure de la Base de Données (texte)
# ------------------------------------------------------------------------------
doc.add_heading("3. Structure de la Base de Données (Texte)", level=2)
metadata = db.Model.metadata
for table in metadata.sorted_tables:
    doc.add_heading(f"Table: {table.name}", level=3)
    for column in table.columns:
        doc.add_paragraph(f"• {column.name} ({column.type})")

# ------------------------------------------------------------------------------
# 4. Diagramme Visuel de la Base de Données
# ------------------------------------------------------------------------------
doc.add_heading("4. Diagramme Visuel de la Base de Données", level=2)

def create_schema_graph(metadata):
    graph = pydot.Dot(graph_type="digraph", rankdir="LR")
    table_nodes = {}
    for table in metadata.tables.values():
        label = f"{{ {table.name} | "
        cols = [f"{col.name} : {col.type}" for col in table.columns]
        label += "\\l".join(cols) + "\\l}"
        node = pydot.Node(table.name, shape="record", label=label, style="filled", fillcolor="lightblue")
        graph.add_node(node)
        table_nodes[table.name] = node
    for table in metadata.tables.values():
        for fk in table.foreign_keys:
            referencing_node = table_nodes[table.name]
            referenced_node = table_nodes[fk.column.table.name]
            edge = pydot.Edge(referencing_node, referenced_node,
                              label=f"{fk.parent.name} → {fk.column.name}", color="blue", arrowsize="0.8")
            graph.add_edge(edge)
    return graph

diagram_path = os.path.join(root_dir, "database_schema.png")
try:
    graph = create_schema_graph(metadata)
    graph.write_png(diagram_path)
    doc.add_paragraph(f"Diagramme généré : {diagram_path}")
except Exception as e:
    doc.add_paragraph(f"❌ Erreur lors de la génération du diagramme : {e}")

# ------------------------------------------------------------------------------
# 5. Sauvegarde finale
# ------------------------------------------------------------------------------
output_path = os.path.join(root_dir, "export_code.docx")
doc.save(output_path)
print(f"✅ Export terminé ! Fichier généré : {output_path}")
